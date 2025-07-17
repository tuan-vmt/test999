#!/usr/bin/env python3
"""
Integrated DOCX Converter with Bookmark Value Preservation
Convert DOCX → PDF → DOCX while preserving all bookmark key-value pairs
"""

import os
import sys
import shutil
import tempfile
import time
from pathlib import Path
import logging

# Import our custom modules
from docx_converter_fixed import DocxToPdfConverter
from bookmark_preserver import BookmarkValuePreserver

# Setup logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

class IntegratedDocxConverter:
    """
    Integrated converter that preserves bookmark values
    Workflow: Extract → Convert → Restore
    """
    
    def __init__(self, keep_intermediates=False):
        self.keep_intermediates = keep_intermediates
        self.temp_dir = tempfile.mkdtemp(prefix='integrated_convert_')
        self.bookmark_preserver = BookmarkValuePreserver()
        self.pdf_converter = None
        
        logger.info(f"Temp directory: {self.temp_dir}")
    
    def convert_with_bookmark_preservation(self, 
                                         input_docx: str, 
                                         output_docx: str,
                                         intermediate_pdf: str = None,
                                         backup_values: str = None,
                                         skip_pdf_conversion: bool = False) -> bool:
        """
        Main conversion method với bookmark preservation
        
        Args:
            input_docx: Input DOCX file
            output_docx: Output DOCX file (with preserved bookmarks)
            intermediate_pdf: Optional path for intermediate PDF
            backup_values: Optional path for bookmark backup JSON
            skip_pdf_conversion: If True, only extract and restore bookmarks without PDF conversion
        
        Returns:
            bool: Success status
        """
        
        print("🚀 === INTEGRATED DOCX CONVERSION WITH BOOKMARK PRESERVATION ===")
        print(f"📁 Input:  {os.path.basename(input_docx)}")
        print(f"📁 Output: {os.path.basename(output_docx)}")
        
        if skip_pdf_conversion:
            print("⚡ Mode: Bookmark-only preservation (skip PDF conversion)")
        else:
            print("🔄 Mode: Full conversion with bookmark preservation")
        print()
        
        try:
            # Step 1: Extract bookmark values
            print("📋 STEP 1: Extracting bookmark values...")
            self.bookmark_preserver.extract_bookmark_values(input_docx)
            
            # Save backup
            if backup_values is None:
                backup_values = os.path.join(self.temp_dir, "bookmark_backup.json")
            
            backup_path = self.bookmark_preserver.save_backup(backup_values)
            if not backup_path:
                logger.error("❌ Failed to save bookmark backup")
                return False
            
            print()
            
            if skip_pdf_conversion:
                # Direct bookmark restoration without PDF conversion
                print("📋 STEP 2: Direct bookmark restoration...")
                
                restore_success = self.bookmark_preserver.restore_bookmark_values(input_docx, output_docx)
                
                if restore_success:
                    print("✅ Bookmark values restored successfully!")
                    return True
                else:
                    logger.error("❌ Failed to restore bookmark values")
                    return False
            
            else:
                # Full conversion workflow
                return self._full_conversion_workflow(input_docx, output_docx, intermediate_pdf, backup_path)
                
        except Exception as e:
            logger.error(f"❌ Conversion failed: {e}")
            return False
        
        finally:
            # Cleanup if not keeping intermediates
            if not self.keep_intermediates:
                self._cleanup()
    
    def _full_conversion_workflow(self, input_docx: str, output_docx: str, 
                                intermediate_pdf: str, backup_path: str) -> bool:
        """Full conversion workflow: DOCX → PDF → DOCX"""
        
        # Step 2: Convert DOCX → PDF
        print("🔄 STEP 2: Converting DOCX to PDF...")
        
        if intermediate_pdf is None:
            intermediate_pdf = os.path.join(self.temp_dir, "intermediate.pdf")
        
        with DocxToPdfConverter() as converter:
            pdf_success = converter.convert_single_file(input_docx, intermediate_pdf)
        
        if not pdf_success:
            logger.error("❌ Failed to convert DOCX to PDF")
            return False
        
        # Verify PDF has bookmarks
        with DocxToPdfConverter() as converter:
            has_bookmarks = converter.verify_pdf_bookmarks(intermediate_pdf, show_details=False)
            if has_bookmarks:
                print("✅ PDF bookmarks preserved successfully")
            else:
                print("⚠️  PDF bookmarks may be missing - will restore from backup")
        
        print()
        
        # Step 3: Convert PDF → DOCX
        print("🔄 STEP 3: Converting PDF back to DOCX...")
        temp_docx = os.path.join(self.temp_dir, "converted.docx")
        
        docx_success = self._convert_pdf_to_docx(intermediate_pdf, temp_docx)
        
        if not docx_success:
            logger.error("❌ Failed to convert PDF to DOCX")
            logger.error("💡 Suggestion: Try --skip-pdf-conversion mode for bookmark-only preservation")
            return False
        
        print()
        
        # Step 4: Restore bookmark values
        print("📋 STEP 4: Restoring bookmark values...")
        
        restore_success = self.bookmark_preserver.restore_bookmark_values(temp_docx, output_docx)
        
        if not restore_success:
            logger.error("❌ Failed to restore bookmark values")
            # Still save the converted file without bookmarks
            shutil.copy2(temp_docx, output_docx)
            print(f"⚠️  Saved converted file without bookmark restoration: {output_docx}")
            return False
        
        print()
        
        # Step 5: Final verification
        print("✅ STEP 5: Final verification...")
        
        if os.path.exists(output_docx):
            file_size = os.path.getsize(output_docx)
            print(f"✅ Output file created: {output_docx} ({file_size:,} bytes)")
            
            # Quick verification of restored content
            self._verify_restored_content(output_docx)
            
            # Summary
            print(f"\n🎉 === CONVERSION COMPLETED SUCCESSFULLY ===")
            print(f"📁 Original: {os.path.basename(input_docx)}")
            print(f"📁 Final:    {os.path.basename(output_docx)}")
            
            if self.keep_intermediates:
                print(f"📁 PDF:      {intermediate_pdf}")
                print(f"📁 Backup:   {backup_path}")
            
            return True
        else:
            logger.error("❌ Output file was not created")
            return False
    
    def _convert_pdf_to_docx(self, pdf_path: str, output_docx: str) -> bool:
        """
        Convert PDF back to DOCX using available tools
        """
        print("🔄 Trying multiple PDF → DOCX conversion methods...")
        
        try:
            # Method 1: Try using pdf2docx (most reliable)
            print("   📝 Method 1: Trying pdf2docx...")
            success = self._convert_pdf_with_pdf2docx(pdf_path, output_docx)
            if success:
                print("   ✅ Success with pdf2docx")
                return True
            else:
                print("   ❌ pdf2docx failed or not available")
            
            # Method 2: Try using PyMuPDF + python-docx
            print("   📝 Method 2: Trying PyMuPDF extraction...")
            success = self._convert_pdf_with_pymupdf(pdf_path, output_docx)
            if success:
                print("   ✅ Success with PyMuPDF")
                return True
            else:
                print("   ❌ PyMuPDF failed or not available")
            
            # Method 3: Try using LibreOffice (reverse conversion)
            print("   📝 Method 3: Trying LibreOffice...")
            success = self._convert_pdf_with_libreoffice(pdf_path, output_docx)
            if success:
                print("   ✅ Success with LibreOffice")
                return True
            else:
                print("   ❌ LibreOffice conversion failed")
            
            # Method 4: Extract text and create DOCX
            print("   📝 Method 4: Trying text extraction...")
            success = self._convert_pdf_with_text_extraction(pdf_path, output_docx)
            if success:
                print("   ✅ Success with text extraction")
                return True
            else:
                print("   ❌ Text extraction failed")
            
            # Fallback: Create reference DOCX
            print("   📝 Fallback: Creating reference document...")
            success = self._create_reference_docx(pdf_path, output_docx)
            if success:
                print("   ⚠️  Created reference DOCX - manual editing needed")
                return True
            
            return False
            
        except Exception as e:
            logger.error(f"PDF to DOCX conversion error: {e}")
            return False
    
    def _convert_pdf_with_libreoffice(self, pdf_path: str, output_docx: str) -> bool:
        """Try converting PDF to DOCX using LibreOffice (improved)"""
        try:
            import subprocess
            
            print(f"      Converting {os.path.basename(pdf_path)} with LibreOffice...")
            
            # Find LibreOffice
            with DocxToPdfConverter() as converter:
                libreoffice_cmd = converter.libreoffice_cmd
                
                if not libreoffice_cmd:
                    return False
                
                # Create temp directory for conversion
                temp_dir = os.path.join(self.temp_dir, "lo_pdf_convert")
                os.makedirs(temp_dir, exist_ok=True)
                
                # Try direct PDF to DOCX conversion
                cmd = [
                    libreoffice_cmd,
                    '--headless',
                    '--invisible',
                    '--nodefault',
                    '--nolockcheck',
                    '--convert-to', 'docx',
                    '--outdir', temp_dir,
                    os.path.abspath(pdf_path)
                ]
                
                logger.debug(f"LibreOffice command: {' '.join(cmd)}")
                
                result = subprocess.run(
                    cmd,
                    capture_output=True,
                    text=True,
                    timeout=120,
                    cwd=temp_dir
                )
                
                if result.returncode == 0:
                    # Look for output file
                    pdf_name = os.path.splitext(os.path.basename(pdf_path))[0]
                    expected_docx = os.path.join(temp_dir, f"{pdf_name}.docx")
                    
                    if os.path.exists(expected_docx) and os.path.getsize(expected_docx) > 1000:
                        shutil.move(expected_docx, output_docx)
                        logger.info("✅ PDF converted to DOCX using LibreOffice")
                        return True
                    else:
                        logger.debug(f"LibreOffice output not found or too small: {expected_docx}")
                        logger.debug(f"Files in temp dir: {os.listdir(temp_dir)}")
                else:
                    logger.debug(f"LibreOffice error (code {result.returncode}): {result.stderr}")
                
                # Alternative: Try PDF → ODT → DOCX
                if not os.path.exists(output_docx):
                    logger.debug("Trying PDF → ODT → DOCX conversion...")
                    
                    # Step 1: PDF → ODT
                    cmd1 = [
                        libreoffice_cmd,
                        '--headless',
                        '--invisible',
                        '--convert-to', 'odt:writer_pdf_import',
                        '--outdir', temp_dir,
                        os.path.abspath(pdf_path)
                    ]
                    
                    result1 = subprocess.run(cmd1, capture_output=True, text=True, timeout=120)
                    
                    if result1.returncode == 0:
                        odt_file = os.path.join(temp_dir, f"{pdf_name}.odt")
                        
                        if os.path.exists(odt_file):
                            # Step 2: ODT → DOCX
                            cmd2 = [
                                libreoffice_cmd,
                                '--headless',
                                '--invisible',
                                '--convert-to', 'docx',
                                '--outdir', temp_dir,
                                odt_file
                            ]
                            
                            result2 = subprocess.run(cmd2, capture_output=True, text=True, timeout=120)
                            
                            if result2.returncode == 0:
                                final_docx = os.path.join(temp_dir, f"{pdf_name}.docx")
                                if os.path.exists(final_docx) and os.path.getsize(final_docx) > 1000:
                                    shutil.move(final_docx, output_docx)
                                    logger.info("✅ PDF converted via ODT using LibreOffice")
                                    return True
            
            return False
            
        except Exception as e:
            logger.debug(f"LibreOffice PDF conversion failed: {e}")
            return False
    
    def _convert_pdf_with_pdf2docx(self, pdf_path: str, output_docx: str) -> bool:
        """Try converting using pdf2docx library (most reliable)"""
        try:
            from pdf2docx import Converter
            
            print(f"      Converting {os.path.basename(pdf_path)} with pdf2docx...")
            
            cv = Converter(pdf_path)
            cv.convert(output_docx, start=0, end=None)
            cv.close()
            
            if os.path.exists(output_docx) and os.path.getsize(output_docx) > 1000:
                logger.info("✅ PDF converted to DOCX using pdf2docx")
                return True
            
            return False
            
        except ImportError:
            logger.debug("pdf2docx not available (install with: pip install pdf2docx)")
            return False
        except Exception as e:
            logger.debug(f"pdf2docx conversion failed: {e}")
            return False
    
    def _convert_pdf_with_pymupdf(self, pdf_path: str, output_docx: str) -> bool:
        """Try converting using PyMuPDF (fitz) for text extraction"""
        try:
            import fitz  # PyMuPDF
            from docx import Document
            from docx.shared import Inches
            
            print(f"      Extracting text from {os.path.basename(pdf_path)} with PyMuPDF...")
            
            # Open PDF
            pdf_doc = fitz.open(pdf_path)
            
            # Create new DOCX
            doc = Document()
            
            # Add title
            doc.add_heading('Converted Document', 0)
            
            # Extract text from each page
            for page_num in range(len(pdf_doc)):
                page = pdf_doc.load_page(page_num)
                
                # Extract text
                text = page.get_text()
                
                if text.strip():
                    # Add page heading
                    if len(pdf_doc) > 1:
                        doc.add_heading(f'Page {page_num + 1}', 1)
                    
                    # Add text content
                    paragraphs = text.split('\n\n')
                    for para_text in paragraphs:
                        para_text = para_text.strip()
                        if para_text:
                            doc.add_paragraph(para_text)
            
            # Close PDF
            pdf_doc.close()
            
            # Save DOCX
            doc.save(output_docx)
            
            if os.path.exists(output_docx) and os.path.getsize(output_docx) > 1000:
                logger.info("✅ PDF converted to DOCX using PyMuPDF text extraction")
                return True
            
            return False
            
        except ImportError:
            logger.debug("PyMuPDF not available (install with: pip install PyMuPDF)")
            return False
        except Exception as e:
            logger.debug(f"PyMuPDF conversion failed: {e}")
            return False
    
    def _convert_pdf_with_text_extraction(self, pdf_path: str, output_docx: str) -> bool:
        """Extract text using multiple methods and create DOCX"""
        try:
            from docx import Document
            import subprocess
            
            print(f"      Extracting text from {os.path.basename(pdf_path)}...")
            
            extracted_text = ""
            
            # Try pdftotext (if available)
            try:
                result = subprocess.run(['pdftotext', pdf_path, '-'], 
                                      capture_output=True, text=True, timeout=30)
                if result.returncode == 0 and result.stdout.strip():
                    extracted_text = result.stdout
                    logger.debug("Text extracted using pdftotext")
            except:
                pass
            
            # Try pdfplumber as fallback
            if not extracted_text:
                try:
                    import pdfplumber
                    
                    with pdfplumber.open(pdf_path) as pdf:
                        text_parts = []
                        for page in pdf.pages:
                            page_text = page.extract_text()
                            if page_text:
                                text_parts.append(page_text)
                        extracted_text = '\n\n'.join(text_parts)
                        logger.debug("Text extracted using pdfplumber")
                except ImportError:
                    logger.debug("pdfplumber not available")
                except Exception as e:
                    logger.debug(f"pdfplumber failed: {e}")
            
            # Create DOCX if we have text
            if extracted_text.strip():
                doc = Document()
                
                # Add title
                doc.add_heading('Converted Document', 0)
                
                # Add note
                note = doc.add_paragraph()
                note.add_run("Note: ").bold = True
                note.add_run(f"Content extracted from: {os.path.basename(pdf_path)}")
                
                # Add content
                doc.add_heading('Document Content', 1)
                
                # Split into paragraphs and add
                paragraphs = extracted_text.split('\n\n')
                for para_text in paragraphs:
                    para_text = para_text.strip()
                    if para_text and len(para_text) > 3:  # Skip very short lines
                        doc.add_paragraph(para_text)
                
                # Save
                doc.save(output_docx)
                
                if os.path.exists(output_docx):
                    logger.info("✅ DOCX created from extracted text")
                    return True
            
            return False
            
        except Exception as e:
            logger.debug(f"Text extraction failed: {e}")
            return False
    
    def _create_reference_docx(self, pdf_path: str, output_docx: str) -> bool:
        """Create a basic DOCX with reference to PDF content"""
        try:
            from docx import Document
            from docx.shared import Inches
            
            doc = Document()
            
            # Add title
            doc.add_heading('Converted Document', 0)
            
            # Add note about conversion
            note = doc.add_paragraph()
            note.add_run("Note: ").bold = True
            note.add_run(f"This document was converted from PDF: {os.path.basename(pdf_path)}. ")
            note.add_run("The bookmark values have been restored from the original DOCX file. ")
            note.add_run("Please review and edit the content as needed.")
            
            # Add a placeholder for content
            doc.add_heading('Document Content', 1)
            content_para = doc.add_paragraph()
            content_para.add_run("Content from PDF conversion will appear here. ")
            content_para.add_run("Please manually copy content from the PDF file or use advanced PDF-to-DOCX tools.")
            
            # Save
            doc.save(output_docx)
            
            logger.info("✅ Reference DOCX created")
            return True
            
        except Exception as e:
            logger.error(f"Failed to create reference DOCX: {e}")
            return False
    
    def _verify_restored_content(self, docx_path: str):
        """Quick verification of restored content"""
        try:
            from docx import Document
            
            doc = Document(docx_path)
            
            # Count paragraphs
            para_count = len([p for p in doc.paragraphs if p.text.strip()])
            
            # Check for reference table (added by bookmark restoration)
            has_reference_table = any("Extracted Field Values" in p.text for p in doc.paragraphs)
            
            print(f"📊 Document verification:")
            print(f"   - Paragraphs: {para_count}")
            print(f"   - Reference table: {'Yes' if has_reference_table else 'No'}")
            
        except Exception as e:
            logger.debug(f"Verification error: {e}")
    
    def _cleanup(self):
        """Clean up temporary files"""
        try:
            if os.path.exists(self.temp_dir):
                shutil.rmtree(self.temp_dir, ignore_errors=True)
                logger.info(f"Temp directory cleaned up: {self.temp_dir}")
        except Exception as e:
            logger.debug(f"Cleanup error: {e}")
    
    def __enter__(self):
        return self
    
    def __exit__(self, exc_type, exc_val, exc_tb):
        if not self.keep_intermediates:
            self._cleanup()

# Convenience functions
def convert_docx_preserving_bookmarks(input_docx: str, 
                                    output_docx: str,
                                    keep_intermediates: bool = False,
                                    skip_pdf_conversion: bool = False) -> bool:
    """
    Simple function để convert DOCX với bookmark preservation
    
    Args:
        input_docx: Input DOCX file
        output_docx: Output DOCX file
        keep_intermediates: Keep intermediate files for debugging
        skip_pdf_conversion: Only restore bookmarks without PDF conversion
    
    Returns:
        bool: Success status
    """
    with IntegratedDocxConverter(keep_intermediates=keep_intermediates) as converter:
        return converter.convert_with_bookmark_preservation(
            input_docx, output_docx, skip_pdf_conversion=skip_pdf_conversion
        )

def restore_bookmarks_only(input_docx: str, output_docx: str) -> bool:
    """
    Only restore bookmark values without any conversion
    """
    return convert_docx_preserving_bookmarks(
        input_docx, output_docx, skip_pdf_conversion=True
    )

def main():
    """Main function"""
    import argparse
    
    parser = argparse.ArgumentParser(
        description='Convert DOCX with bookmark value preservation',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Simple bookmark-only preservation (recommended)
  python integrated_converter.py input.docx output.docx --skip-pdf-conversion
  
  # Full conversion with PDF (may have formatting issues)
  python integrated_converter.py input.docx output.docx --keep-intermediates
  
  # Install PDF conversion dependencies
  python integrated_converter.py --install-deps
  
  # Specify custom paths
  python integrated_converter.py input.docx output.docx --pdf temp.pdf --backup values.json
        """
    )
    
    parser.add_argument('input', help='Input DOCX file')
    parser.add_argument('output', help='Output DOCX file')
    parser.add_argument('--pdf', help='Intermediate PDF file path')
    parser.add_argument('--backup', help='Bookmark backup JSON file path')
    parser.add_argument('--keep-intermediates', action='store_true', 
                       help='Keep intermediate files (PDF, backup) for debugging')
    parser.add_argument('--skip-pdf-conversion', action='store_true',
                       help='Only restore bookmarks without PDF conversion (faster, preserves formatting)')
    parser.add_argument('--install-deps', action='store_true',
                       help='Install required dependencies for PDF conversion')
    
    args = parser.parse_args()
    
    # Handle dependency installation
    if args.install_deps:
        print("📦 Installing PDF conversion dependencies...")
        try:
            import subprocess
            import sys
            
            dependencies = ['pdf2docx', 'PyMuPDF', 'pdfplumber']
            for dep in dependencies:
                print(f"   Installing {dep}...")
                subprocess.check_call([sys.executable, '-m', 'pip', 'install', dep])
            
            print("✅ Dependencies installed successfully!")
            return 0
            
        except Exception as e:
            print(f"❌ Failed to install dependencies: {e}")
            return 1
    
    # Validate input
    if not os.path.exists(args.input):
        print(f"❌ Input file not found: {args.input}")
        return 1
    
    if not args.input.lower().endswith('.docx'):
        print(f"❌ Input must be a DOCX file: {args.input}")
        return 1
    
    # Ensure output directory exists
    output_dir = os.path.dirname(args.output)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir, exist_ok=True)
    
    try:
        # Run conversion
        with IntegratedDocxConverter(keep_intermediates=args.keep_intermediates) as converter:
            success = converter.convert_with_bookmark_preservation(
                args.input, 
                args.output,
                intermediate_pdf=args.pdf,
                backup_values=args.backup,
                skip_pdf_conversion=args.skip_pdf_conversion
            )
        
        return 0 if success else 1
        
    except KeyboardInterrupt:
        print("\n❌ Conversion cancelled by user")
        return 1
    except Exception as e:
        print(f"❌ Conversion error: {e}")
        return 1

if __name__ == "__main__":
    import subprocess  # Add missing import
    sys.exit(main())