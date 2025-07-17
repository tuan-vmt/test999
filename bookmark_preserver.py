#!/usr/bin/env python3
"""
DOCX Bookmark Value Preserver
Extract bookmark key-values trước convert, restore sau convert
"""

import os
import sys
import json
import shutil
import tempfile
from pathlib import Path
from typing import Dict, List, Any, Optional
import zipfile
import xml.etree.ElementTree as ET
from datetime import datetime

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Inches
except ImportError:
    print("❌ Cần cài đặt: pip install python-docx")
    sys.exit(1)

class BookmarkValuePreserver:
    """Class để preserve bookmark values qua quá trình convert"""
    
    def __init__(self):
        self.original_values = {}
        self.backup_file = None
        
    def extract_bookmark_values(self, docx_path: str) -> Dict[str, Any]:
        """Extract tất cả bookmark values từ DOCX gốc"""
        print(f"🔍 Extracting bookmark values từ: {os.path.basename(docx_path)}")
        
        extracted_data = {
            'file_path': docx_path,
            'extraction_time': datetime.now().isoformat(),
            'bookmarks': {},
            'content_controls': {},
            'empty_fields': {},
            'form_fields': {},
            'placeholder_patterns': {}
        }
        
        # Extract bookmarks với XML parsing
        self._extract_bookmarks_xml(docx_path, extracted_data)
        
        # Extract content controls
        self._extract_content_controls_xml(docx_path, extracted_data)
        
        # Extract các field patterns từ text
        self._extract_field_patterns(docx_path, extracted_data)
        
        # Extract empty fields (có thể là bookmarks bị mất value)
        self._extract_empty_fields(docx_path, extracted_data)
        
        total_fields = (len(extracted_data['bookmarks']) + 
                       len(extracted_data['content_controls']) +
                       len(extracted_data['empty_fields']) +
                       len(extracted_data['placeholder_patterns']))
        
        print(f"✅ Extracted {total_fields} fields:")
        print(f"   - Bookmarks: {len(extracted_data['bookmarks'])}")
        print(f"   - Content Controls: {len(extracted_data['content_controls'])}")
        print(f"   - Empty Fields: {len(extracted_data['empty_fields'])}")
        print(f"   - Patterns: {len(extracted_data['placeholder_patterns'])}")
        
        self.original_values = extracted_data
        return extracted_data
    
    def _extract_bookmarks_xml(self, docx_path: str, data: Dict):
        """Extract bookmarks bằng XML parsing"""
        try:
            with zipfile.ZipFile(docx_path, 'r') as zip_file:
                doc_xml = zip_file.read('word/document.xml').decode('utf-8')
                root = ET.fromstring(doc_xml.encode('utf-8'))
                
                namespaces = {
                    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                }
                
                # Find all bookmark starts
                for bookmark_start in root.findall('.//w:bookmarkStart', namespaces):
                    try:
                        bookmark_id = bookmark_start.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
                        bookmark_name = bookmark_start.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}name')
                        
                        if bookmark_name and not bookmark_name.startswith('_'):
                            # Get content between bookmark start and end
                            content = self._get_bookmark_content_xml(root, bookmark_id, namespaces)
                            
                            data['bookmarks'][bookmark_name] = {
                                'name': bookmark_name,
                                'value': content,
                                'id': bookmark_id,
                                'type': 'bookmark'
                            }
                            
                    except Exception as e:
                        print(f"⚠️  Error parsing bookmark: {e}")
                        continue
                        
        except Exception as e:
            print(f"⚠️  Error extracting bookmarks: {e}")
    
    def _get_bookmark_content_xml(self, root, bookmark_id: str, namespaces: dict) -> str:
        """Get content between bookmark start and end"""
        try:
            content_parts = []
            in_bookmark = False
            
            # Walk through all elements in document order
            for elem in root.iter():
                # Check for bookmark start
                if (elem.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bookmarkStart' and 
                    elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id') == bookmark_id):
                    in_bookmark = True
                    continue
                
                # Check for bookmark end
                if (elem.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bookmarkEnd' and 
                    elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id') == bookmark_id):
                    break
                
                # Collect text content within bookmark
                if in_bookmark and elem.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t':
                    if elem.text:
                        content_parts.append(elem.text)
            
            return ''.join(content_parts).strip()
            
        except Exception:
            return ""
    
    def _extract_content_controls_xml(self, docx_path: str, data: Dict):
        """Extract content controls"""
        try:
            with zipfile.ZipFile(docx_path, 'r') as zip_file:
                doc_xml = zip_file.read('word/document.xml').decode('utf-8')
                root = ET.fromstring(doc_xml.encode('utf-8'))
                
                namespaces = {
                    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                }
                
                for sdt in root.findall('.//w:sdt', namespaces):
                    try:
                        sdt_pr = sdt.find('w:sdtPr', namespaces)
                        if sdt_pr is None:
                            continue
                        
                        # Get tag (key)
                        tag_elem = sdt_pr.find('w:tag', namespaces)
                        tag_val = tag_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') if tag_elem is not None else None
                        
                        # Get alias (display name)
                        alias_elem = sdt_pr.find('w:alias', namespaces)
                        alias_val = alias_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') if alias_elem is not None else None
                        
                        # Get current content
                        sdt_content = sdt.find('w:sdtContent', namespaces)
                        content_text = ""
                        if sdt_content is not None:
                            text_elems = sdt_content.findall('.//w:t', namespaces)
                            content_text = ''.join([t.text or '' for t in text_elems])
                        
                        # Create unique key
                        field_key = tag_val or alias_val or f"content_control_{len(data['content_controls']) + 1}"
                        
                        data['content_controls'][field_key] = {
                            'name': field_key,
                            'value': content_text.strip(),
                            'tag': tag_val,
                            'alias': alias_val,
                            'type': 'content_control'
                        }
                        
                    except Exception as e:
                        print(f"⚠️  Error parsing content control: {e}")
                        continue
                        
        except Exception as e:
            print(f"⚠️  Error extracting content controls: {e}")
    
    def _extract_field_patterns(self, docx_path: str, data: Dict):
        """Extract placeholder patterns như [text], {text}"""
        try:
            doc = Document(docx_path)
            full_text = '\n'.join([para.text for para in doc.paragraphs])
            
            import re
            patterns = [
                (r'\[([^\]]+)\]', 'bracket'),
                (r'\{([^}]+)\}', 'brace'),
                (r'<<([^>]+)>>', 'double_angle'),
                (r'{{([^}]+)}}', 'double_brace'),
                (r'\$\{([^}]+)\}', 'dollar_brace'),
            ]
            
            for pattern, pattern_type in patterns:
                matches = re.finditer(pattern, full_text)
                for match in matches:
                    placeholder = match.group(1).strip()
                    full_match = match.group(0)
                    
                    if 2 <= len(placeholder) <= 100:
                        key = f"{pattern_type}_{placeholder.replace(' ', '_')}"
                        
                        data['placeholder_patterns'][key] = {
                            'name': placeholder,
                            'value': full_match,
                            'pattern_type': pattern_type,
                            'type': 'placeholder'
                        }
            
        except Exception as e:
            print(f"⚠️  Error extracting patterns: {e}")
    
    def _extract_empty_fields(self, docx_path: str, data: Dict):
        """Extract empty fields (có thể là bookmarks bị mất value)"""
        try:
            doc = Document(docx_path)
            full_text = '\n'.join([para.text for para in doc.paragraphs])
            
            import re
            patterns = [
                (r'(\w+):\s{3,}', 'label_with_spaces'),
                (r'(\w+\s+\w+):\s{3,}', 'two_word_label'),
                (r'([A-Za-zÀ-ỹ\s]+):\s{3,}', 'vietnamese_label'),
            ]
            
            field_counter = 1
            for pattern, pattern_type in patterns:
                matches = re.finditer(pattern, full_text, re.MULTILINE)
                for match in matches:
                    if len(match.groups()) >= 1:
                        label = match.group(-1).strip()
                        
                        if label and len(label) > 1:
                            key = f"empty_field_{field_counter}_{label.replace(' ', '_')}"
                            
                            data['empty_fields'][key] = {
                                'name': label,
                                'value': '',  # Empty field
                                'pattern_type': pattern_type,
                                'type': 'empty_field'
                            }
                            field_counter += 1
            
        except Exception as e:
            print(f"⚠️  Error extracting empty fields: {e}")
    
    def save_backup(self, output_path: str = None) -> str:
        """Save extracted values to backup file"""
        if not self.original_values:
            raise ValueError("No extracted values to save. Run extract_bookmark_values first.")
        
        if output_path is None:
            base_name = os.path.splitext(self.original_values['file_path'])[0]
            output_path = f"{base_name}_bookmark_backup.json"
        
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(self.original_values, f, ensure_ascii=False, indent=2)
            
            self.backup_file = output_path
            print(f"💾 Backup saved: {output_path}")
            return output_path
        except Exception as e:
            print(f"❌ Error saving backup: {e}")
            return ""
    
    def load_backup(self, backup_path: str) -> bool:
        """Load backup values from file"""
        try:
            with open(backup_path, 'r', encoding='utf-8') as f:
                self.original_values = json.load(f)
            
            self.backup_file = backup_path
            print(f"📂 Backup loaded: {backup_path}")
            return True
        except Exception as e:
            print(f"❌ Error loading backup: {e}")
            return False
    
    def restore_bookmark_values(self, target_docx_path: str, output_path: str = None) -> bool:
        """Restore bookmark values vào DOCX mới"""
        if not self.original_values:
            print("❌ No backup values available. Load backup first.")
            return False
        
        if output_path is None:
            base_name = os.path.splitext(target_docx_path)[0]
            output_path = f"{base_name}_restored.docx"
        
        print(f"🔄 Restoring values to: {os.path.basename(output_path)}")
        
        try:
            # Copy target file to output
            shutil.copy2(target_docx_path, output_path)
            
            # Open document for editing
            doc = Document(output_path)
            
            restored_count = 0
            
            # Method 1: Try to restore through content replacement
            restored_count += self._restore_through_text_replacement(doc)
            
            # Method 2: Try to find and restore bookmarks
            restored_count += self._restore_through_bookmark_insertion(doc)
            
            # Method 3: Add missing values as comments or at end
            restored_count += self._add_missing_values_as_reference(doc)
            
            # Save document
            doc.save(output_path)
            
            print(f"✅ Restored {restored_count} values to: {output_path}")
            return True
            
        except Exception as e:
            print(f"❌ Error restoring values: {e}")
            return False
    
    def _restore_through_text_replacement(self, doc: Document) -> int:
        """Restore values bằng cách thay thế text patterns"""
        restored = 0
        
        try:
            # Get all values to restore
            all_values = {}
            
            # Combine all field types
            for bookmark_name, info in self.original_values.get('bookmarks', {}).items():
                if info['value'].strip():  # Only restore non-empty values
                    all_values[bookmark_name] = info['value']
            
            for cc_name, info in self.original_values.get('content_controls', {}).items():
                if info['value'].strip():
                    all_values[cc_name] = info['value']
            
            for pattern_name, info in self.original_values.get('placeholder_patterns', {}).items():
                if info['value'].strip():
                    all_values[info['name']] = info['value']
            
            # Try to replace empty patterns with values
            for paragraph in doc.paragraphs:
                original_text = paragraph.text
                updated_text = original_text
                
                # Look for empty fields and fill them
                import re
                
                # Pattern: "Label:     " -> "Label: value"
                for field_name, field_value in all_values.items():
                    patterns = [
                        f"{field_name}:\\s{{3,}}",  # "field_name:   "
                        f"{field_name}\\s*:\\s{{3,}}",  # "field_name :   "
                        f"\\b{field_name}\\b\\s*:\\s*$",  # "field_name: " at end of line
                    ]
                    
                    for pattern in patterns:
                        if re.search(pattern, updated_text, re.IGNORECASE):
                            replacement = f"{field_name}: {field_value}"
                            updated_text = re.sub(pattern, replacement, updated_text, flags=re.IGNORECASE)
                            restored += 1
                            print(f"   ✓ Restored '{field_name}': '{field_value}'")
                            break
                
                # Update paragraph if changed
                if updated_text != original_text:
                    paragraph.text = updated_text
            
        except Exception as e:
            print(f"⚠️  Error in text replacement: {e}")
        
        return restored
    
    def _restore_through_bookmark_insertion(self, doc: Document) -> int:
        """Try to restore through bookmark insertion (limited support in python-docx)"""
        restored = 0
        
        # Note: python-docx has limited bookmark support
        # This is a placeholder for more advanced bookmark manipulation
        
        return restored
    
    def _add_missing_values_as_reference(self, doc: Document) -> int:
        """Add missing values as reference table at end of document"""
        try:
            # Collect all values that should be restored
            missing_values = {}
            
            for bookmark_name, info in self.original_values.get('bookmarks', {}).items():
                if info['value'].strip():
                    missing_values[bookmark_name] = info['value']
            
            for cc_name, info in self.original_values.get('content_controls', {}).items():
                if info['value'].strip():
                    missing_values[cc_name] = info['value']
            
            if missing_values:
                # Add a section break
                doc.add_page_break()
                
                # Add header
                header = doc.add_heading('📋 Extracted Field Values (Reference)', 2)
                
                # Add values as table
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Light Shading Accent 1'
                
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Field Name'
                hdr_cells[1].text = 'Value'
                
                for field_name, field_value in missing_values.items():
                    row_cells = table.add_row().cells
                    row_cells[0].text = field_name
                    row_cells[1].text = str(field_value)[:200]  # Limit length
                
                # Add note
                note = doc.add_paragraph()
                note.add_run("Note: ").bold = True
                note.add_run("These are the original field values extracted before conversion. ")
                note.add_run("Use this reference to manually restore values if needed.")
                
                return len(missing_values)
        
        except Exception as e:
            print(f"⚠️  Error adding reference table: {e}")
        
        return 0
    
    def print_summary(self):
        """Print summary of extracted values"""
        if not self.original_values:
            print("❌ No extracted values available")
            return
        
        data = self.original_values
        print(f"\n📊 === BOOKMARK VALUES SUMMARY ===")
        print(f"📁 Original file: {os.path.basename(data['file_path'])}")
        print(f"⏰ Extracted: {data['extraction_time']}")
        
        if data.get('bookmarks'):
            print(f"\n🔖 Bookmarks ({len(data['bookmarks'])}):")
            for name, info in data['bookmarks'].items():
                value_preview = info['value'][:50] + "..." if len(info['value']) > 50 else info['value']
                print(f"   • {name}: '{value_preview}'")
        
        if data.get('content_controls'):
            print(f"\n📋 Content Controls ({len(data['content_controls'])}):")
            for name, info in data['content_controls'].items():
                value_preview = info['value'][:50] + "..." if len(info['value']) > 50 else info['value']
                print(f"   • {name}: '{value_preview}'")
        
        if data.get('empty_fields'):
            print(f"\n📝 Empty Fields ({len(data['empty_fields'])}):")
            for name, info in data['empty_fields'].items():
                print(f"   • {info['name']}: [EMPTY FIELD]")

# Convenience functions
def extract_and_backup_values(docx_path: str, backup_path: str = None) -> str:
    """Extract bookmark values và save backup"""
    preserver = BookmarkValuePreserver()
    preserver.extract_bookmark_values(docx_path)
    return preserver.save_backup(backup_path)

def restore_values_from_backup(backup_path: str, target_docx: str, output_path: str = None) -> bool:
    """Restore values từ backup file"""
    preserver = BookmarkValuePreserver()
    if preserver.load_backup(backup_path):
        return preserver.restore_bookmark_values(target_docx, output_path)
    return False

def main():
    """Main function"""
    import argparse
    
    parser = argparse.ArgumentParser(
        description='Preserve bookmark values through DOCX conversion',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Extract and backup values
  python preserve_bookmarks.py extract input.docx
  
  # Restore values to converted document
  python preserve_bookmarks.py restore backup.json converted.docx
  
  # Full workflow: extract -> backup -> restore
  python preserve_bookmarks.py extract original.docx -o values.json
  python preserve_bookmarks.py restore values.json converted.docx -o final.docx
        """
    )
    
    subparsers = parser.add_subparsers(dest='command', help='Commands')
    
    # Extract command
    extract_parser = subparsers.add_parser('extract', help='Extract bookmark values')
    extract_parser.add_argument('input', help='Input DOCX file')
    extract_parser.add_argument('-o', '--output', help='Output backup JSON file')
    
    # Restore command  
    restore_parser = subparsers.add_parser('restore', help='Restore bookmark values')
    restore_parser.add_argument('backup', help='Backup JSON file')
    restore_parser.add_argument('target', help='Target DOCX file to restore to')
    restore_parser.add_argument('-o', '--output', help='Output DOCX file')
    
    args = parser.parse_args()
    
    if args.command == 'extract':
        if not os.path.exists(args.input):
            print(f"❌ File not found: {args.input}")
            return 1
        
        try:
            preserver = BookmarkValuePreserver()
            preserver.extract_bookmark_values(args.input)
            preserver.print_summary()
            backup_path = preserver.save_backup(args.output)
            
            if backup_path:
                print(f"\n✅ Values backed up to: {backup_path}")
                return 0
            else:
                return 1
                
        except Exception as e:
            print(f"❌ Error: {e}")
            return 1
    
    elif args.command == 'restore':
        if not os.path.exists(args.backup):
            print(f"❌ Backup file not found: {args.backup}")
            return 1
        
        if not os.path.exists(args.target):
            print(f"❌ Target file not found: {args.target}")
            return 1
        
        try:
            preserver = BookmarkValuePreserver()
            if preserver.load_backup(args.backup):
                preserver.print_summary()
                success = preserver.restore_bookmark_values(args.target, args.output)
                return 0 if success else 1
            else:
                return 1
                
        except Exception as e:
            print(f"❌ Error: {e}")
            return 1
    
    else:
        parser.print_help()
        return 1

if __name__ == "__main__":
    sys.exit(main())